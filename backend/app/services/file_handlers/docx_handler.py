import logging
from docx import Document
import re
import uuid
from datetime import datetime
from io import BytesIO
import traceback
import os
from openai import AzureOpenAI
import json
from requests.exceptions import Timeout

# Configure logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Helper function to create a client
def create_client(api_key, api_version, endpoint):
    logger.debug(f"Creating AzureOpenAI client with endpoint: {endpoint}")
    return AzureOpenAI(
        api_key=api_key,
        api_version=api_version,
        azure_endpoint=endpoint
    )

# Load primary and fallback configurations
PRIMARY_CONFIG = {
    "api_key": os.getenv("AZURE_OPENAI_API_KEY"),
    "api_version": os.getenv("AZURE_OPENAI_API_VERSION"),
    "endpoint": os.getenv("AZURE_OPENAI_ENDPOINT"),
    "deployment_name": os.getenv("AZURE_OPENAI_CHAT_DEPLOYMENT_NAME")
}

FALLBACK_CONFIGS = [
    {
        "api_key": os.getenv("AZURE_OPENAI_API_KEY_KOREA"),
        "api_version": os.getenv("AZURE_OPENAI_API_VERSION"),
        "endpoint": os.getenv("AZURE_OPENAI_ENDPOINT_KOREA"),
        "deployment_name": os.getenv("AZURE_OPENAI_CHAT_DEPLOYMENT_NAME_KOREA_1")
    },
    {
        "api_key": os.getenv("AZURE_OPENAI_API_KEY_KOREA"),
        "api_version": os.getenv("AZURE_OPENAI_API_VERSION"),
        "endpoint": os.getenv("AZURE_OPENAI_ENDPOINT_KOREA"),
        "deployment_name": os.getenv("AZURE_OPENAI_CHAT_DEPLOYMENT_NAME_KOREA_2")
    },
    {
        "api_key": os.getenv("AZURE_OPENAI_API_KEY_AUSTRALIA"),
        "api_version": os.getenv("AZURE_OPENAI_API_VERSION"),
        "endpoint": os.getenv("AZURE_OPENAI_ENDPOINT_AUSTRALIA"),
        "deployment_name": os.getenv("AZURE_OPENAI_CHAT_DEPLOYMENT_NAME_AUSTRALIA")
    }
]

logger.info("Initializing primary client")
primary_client = create_client(
    PRIMARY_CONFIG["api_key"],
    PRIMARY_CONFIG["api_version"],
    PRIMARY_CONFIG["endpoint"]
)

def timestamp_to_seconds(timestamp):
    """Convert MM:SS timestamp to seconds."""
    try:
        if not timestamp:
            return 0
        parts = timestamp.split(':')
        if len(parts) == 2:
            minutes, seconds = map(int, parts)
            return minutes * 60 + seconds
        return 0
    except (ValueError, AttributeError) as e:
        logger.error(f"Error converting timestamp {timestamp}: {e}")
        return 0

def is_chinese_char(char):
    """Check if a character is Chinese."""
    ranges = [
        ('\u4e00', '\u9fff'),     # CJK Unified Ideographs
        ('\u3400', '\u4dbf'),     # CJK Unified Ideographs Extension A
        ('\u20000', '\u2a6df'),   # CJK Unified Ideographs Extension B
        ('\u2a700', '\u2b73f'),   # CJK Unified Ideographs Extension C
        ('\u2b740', '\u2b81f'),   # CJK Unified Ideographs Extension D
        ('\u2b820', '\u2ceaf'),   # CJK Unified Ideographs Extension E
        ('\uf900', '\ufaff'),     # CJK Compatibility Ideographs
    ]
    return any(start <= char <= end for start, end in ranges)

def split_into_words(text):
    """Split text into words, treating each Chinese character as a separate word."""
    words = []
    current_word = ""
    for char in text:
        if is_chinese_char(char):
            if current_word:
                words.append(current_word)
                current_word = ""
            words.append(char)
        elif char.isspace():
            if current_word:
                words.append(current_word)
                current_word = ""
        else:
            current_word += char
    if current_word:
        words.append(current_word)
    return words

def clean_text(text):
    """Remove markdown-style formatting and clean text."""
    # Remove ** markers
    text = re.sub(r'\*\*|\*', '', text)
    # Remove timestamps
    text = re.sub(r'\d{2}:\d{2}', '', text)
    # Clean extra whitespace
    text = ' '.join(text.split())
    return text.strip()

def hyphenate_speaker_name(name):
    """
    Replace spaces in speaker names with hyphens.
    Also handles cases with colons and other potential formatting.
    """
    if not name:
        return name
    
    # Remove trailing colon if present
    name = name.rstrip(':')
    
    # Replace spaces with hyphens
    return name.replace(' ', '-')

def extract_speaker_and_text(line):
    """Extract speaker and text from a line."""
    speaker_match = re.match(r'\*\*([^*]+)\*\*:?\s*(.*)', line)
    if speaker_match:
        speaker = hyphenate_speaker_name(speaker_match.group(1).strip())
        text = speaker_match.group(2).strip()
        return speaker, text
    return None, None

def extract_timestamp(text):
    """Extract timestamp from text."""
    timestamp_match = re.search(r'(\d{2}:\d{2})', text)
    if timestamp_match:
        return timestamp_match.group(1)
    return None

def get_extraction_function(first_page_content):
    logger.info("Generating extraction function from first page content")
    prompt = f"""
    Analyze the following content from the first page of a document and provide a Python function to extract speaker, timestamp, and text:

    {first_page_content}

    Here's an example of a function that extracts speaker and text (but not timestamp) from a similar format:

    def extract_speaker_and_text(line):
        \"\"\"Extract speaker and text from a line.\"\"\"
        speaker_match = re.match(r'\\*\\*([^*]+)\\*\\*:?\\s*(.*)', line)
        if speaker_match:
            speaker = hyphenate_speaker_name(speaker_match.group(1).strip())
            text = speaker_match.group(2).strip()
            return speaker, text
        return None, None

    This function was designed for content like:

    **Pre-interview Conversation:**

    **Jing Boran:** 00:00
    Yeah, no one listens to these things. Really.
    **Interviewer:** 00:04
    We actually expect more personal stories or those related to the product. For example, would you talk about some jewelry? Because I've seen you have some home collections or personal experiences to share with us. No rush, let's adjust the position first. Lighting okay, good. Are you happy? We've started.
    **Interviewer:** 00:43
    Who's shouting?

    Create a Python function named 'extract_speaker_and_text' that takes a single 'line' parameter and returns a tuple of (speaker, timestamp, text).

    Requirements:
    1. Handle multiple formats, including:
       a. Chinese format: '说话人1 00:00 Text'
       b. English format with speaker names: 'Interviewer: 01:01 Text'
       c. English format with asterisks: '**Speaker:** 00:00 Text'
       d. Format without timestamps: '**Speaker:** Text' or 'Speaker: Text'
    2. If no speaker is found, return (None, None, line.strip()).
    3. If no timestamp is present, return None for the timestamp.
    4. Use regular expressions (re module) for pattern matching.
    5. Implement the function to be robust and handle variations in spacing and formatting.
    6. Ensure the function cleans up any extra whitespace or special characters.
    7. The function should prioritize matching the most specific patterns first.
    8. Include the 'import re' statement at the beginning of the function.

    Provide only the Python code for the function, without any explanations. The function should be efficient and handle all the cases present in the given content.
    """

    system_content = "You are a helpful assistant designed to output JSON. Provide the Python function as a string value for the 'function' key in the JSON response. The function should handle various formats of speaker identification and timestamps."
    user_content = prompt

    logger.debug(f"Creating AzureOpenAI client for function generation")
    client = AzureOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        api_version="2024-03-01-preview"
    )

    def call_azure_openai(deployment_name):
        logger.info(f"Calling Azure OpenAI with deployment: {deployment_name}")
        try:
            response = client.chat.completions.create(
                model=deployment_name,
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": system_content},
                    {"role": "user", "content": user_content}
                ],
                timeout=180
            )
            logger.debug(f"Received response from Azure OpenAI: {response.choices[0].message.content}")
            return json.loads(response.choices[0].message.content), None
        except json.JSONDecodeError as e:
            logger.error(f"Error decoding JSON response: {str(e)}")
            return None, f"JSON decode error: {str(e)}"
        except Exception as e:
            logger.error(f"Error during API call: {str(e)}")
            return None, str(e)

    # Try primary deployment
    logger.info("Attempting to use primary deployment")
    result, error = call_azure_openai(os.getenv("AZURE_OPENAI_CHAT_DEPLOYMENT_NAME"))
    if result and 'function' in result:
        logger.info("Successfully generated function using primary deployment")
        
        # Clean up the indentation
        function_code = result['function'].strip()
        lines = function_code.split('\n')
        cleaned_lines = [line.strip() for line in lines]
        cleaned_function = '\n'.join(cleaned_lines)
        
        logger.debug(f"Cleaned extraction function code:\n{cleaned_function}")
        return cleaned_function

    # If primary deployment fails, try fallbacks
    fallback_deployments = [
        os.getenv("AZURE_OPENAI_CHAT_DEPLOYMENT_NAME_KOREA_1"),
        os.getenv("AZURE_OPENAI_CHAT_DEPLOYMENT_NAME_KOREA_2"),
        os.getenv("AZURE_OPENAI_CHAT_DEPLOYMENT_NAME_AUSTRALIA")
    ]

    for deployment in fallback_deployments:
        if deployment:
            logger.info(f"Switching to fallback deployment: {deployment}")
            result, error = call_azure_openai(deployment)
            if result and 'function' in result:
                logger.info(f"Successfully generated function using fallback deployment: {deployment}")
                return result['function']

    logger.error("All API calls failed to generate extraction function")
    raise Exception("All API calls failed to generate extraction function")

async def parse_to_schema(file_or_content):
    """Parse DOCX interview transcript to structured schema."""
    try:
        logger.info("Starting parse_to_schema function")
        if isinstance(file_or_content, str):
            logger.debug("Reading file from disk")
            with open(file_or_content, 'rb') as f:
                file_content = f.read()
            filename = os.path.basename(file_or_content)
        elif isinstance(file_or_content, bytes):
            logger.debug("Using provided bytes as file content")
            file_content = file_or_content
            filename = "unknown.docx"
        else:
            logger.debug("Reading file content from uploaded file")
            file_content = await file_or_content.read()
            filename = file_or_content.filename

        logger.info(f"Processing file: {filename}")
        doc = Document(BytesIO(file_content))
        
        logger.debug("Extracting first page content")
        first_page_content = "\n".join([p.text for p in doc.paragraphs[:10]])
        
        logger.info("Getting extraction function from Azure OpenAI")
        extraction_function_code = get_extraction_function(first_page_content)
        
        logger.debug(f"Received extraction function code:\n{extraction_function_code}")

        logger.info("Executing generated extraction function")
        try:
            exec(extraction_function_code, globals())
        except SyntaxError as e:
            logger.error(f"Syntax error in generated code: {e}")
            logger.error(f"Generated code:\n{extraction_function_code}")
            raise
        except Exception as e:
            logger.error(f"Error executing generated code: {e}")
            logger.error(f"Generated code:\n{extraction_function_code}")
            raise
        
        logger.info("Parsing document content")
        segments = []
        project_id = str(uuid.uuid4())
        media_id = str(uuid.uuid4())
        uploaded_on = datetime.utcnow().isoformat() + "Z"

        current_segment = None
        segment_index = 1

        for para in doc.paragraphs:
            text = para.text.strip()
            logger.debug(f"Processing paragraph: {text}")
            if not text:
                continue

            speaker, remaining_text = extract_speaker_and_text(text)
            timestamp = extract_timestamp(text)

            if speaker:
                logger.debug(f"Found new speaker: {speaker}")
                if current_segment:
                    segments.append(current_segment)
                
                start_time = timestamp_to_seconds(timestamp) if timestamp else 0
                words = split_into_words(clean_text(remaining_text))
                current_segment = {
                    "index": segment_index,
                    "start_time": start_time,
                    "end_time": start_time + 60,  # Default duration of 60 seconds
                    "text": clean_text(remaining_text),
                    "speaker": speaker,
                    "words": [{"start": -1, "end": -1, "word": word} for word in words]
                }
                segment_index += 1
            elif current_segment:
                logger.debug("Adding text to current segment")
                cleaned_text = clean_text(text)
                current_segment["text"] += " " + cleaned_text
                words = split_into_words(cleaned_text)
                current_segment["words"].extend([{"start": -1, "end": -1, "word": word} for word in words])

        if current_segment:
            segments.append(current_segment)

        logger.debug("Adjusting segment end times")
        for i in range(len(segments) - 1):
            segments[i]["end_time"] = segments[i+1]["start_time"]

        total_duration = segments[-1]["end_time"] if segments else 0

        parsed_data = {
            "project_id": project_id,
            "media": {
                "id": media_id,
                "source": filename,
                "duration": total_duration,
                "uploaded_on": uploaded_on
            },
            "transcript": {
                "segments": segments
            },
            "edits": []
        }

        logger.info(f"Successfully parsed {len(segments)} segments from DOCX")
        return parsed_data

    except Exception as e:
        logger.error(f"Error parsing DOCX file: {str(e)}")
        logger.error(traceback.format_exc())
        raise